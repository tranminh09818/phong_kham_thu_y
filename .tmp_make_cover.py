from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION_START
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image, ImageDraw, ImageFont
from pathlib import Path
import shutil, math

root = Path(r'D:\QLy Phòng Khám Thú Y')
out = root / 'Báo cáo thực tập cnPKTY.docx'
backup = root / 'Báo cáo thực tập cnPKTY - bản trắng backup.docx'
logo = root / '.tmp_vnua_logo.png'
if out.exists() and not backup.exists():
    shutil.copy2(out, backup)

# Draw an approximate VNUA-style emblem so the cover visually follows the provided sample.
size = 900
img = Image.new('RGBA', (size, size), (255, 255, 255, 0))
d = ImageDraw.Draw(img)
c = size // 2
r = 360
d.ellipse((c-r, c-r, c+r, c+r), fill=(0, 132, 64, 255), outline=(255, 224, 0, 255), width=28)
d.ellipse((c-240, c-240, c+240, c+240), outline=(255, 224, 0, 255), width=12)
try:
    font_ring = ImageFont.truetype('arialbd.ttf', 54)
    font_year = ImageFont.truetype('arialbd.ttf', 76)
except Exception:
    font_ring = ImageFont.load_default()
    font_year = ImageFont.load_default()
# simple wheat petals
for i in range(10):
    ang = math.radians(i * 36 - 90)
    x = c + math.cos(ang) * 155
    y = c + math.sin(ang) * 155
    petal = [(x + math.cos(ang)*85, y + math.sin(ang)*85),
             (x + math.cos(ang+1.85)*38, y + math.sin(ang+1.85)*38),
             (x - math.cos(ang)*20, y - math.sin(ang)*20),
             (x + math.cos(ang-1.85)*38, y + math.sin(ang-1.85)*38)]
    d.polygon(petal, fill=(255, 168, 22, 255), outline=(255, 224, 0, 255))
# ring text approximated as straight blocks to avoid fragile curved text rendering
for text, y in [('HOC VIEN NONG NGHIEP VIET NAM', 215), ('* 1956 *', 615)]:
    font = font_ring if 'HOC' in text else font_year
    bbox = d.textbbox((0,0), text, font=font)
    d.text((c - (bbox[2]-bbox[0])/2, y), text, fill=(255,224,0,255), font=font)
img.save(logo)

doc = Document()
sec = doc.sections[0]
sec.top_margin = Cm(1.5)
sec.bottom_margin = Cm(1.4)
sec.left_margin = Cm(2.2)
sec.right_margin = Cm(2.2)

# page border
sectPr = sec._sectPr
pgBorders = OxmlElement('w:pgBorders')
pgBorders.set(qn('w:offsetFrom'), 'page')
for side in ['top','left','bottom','right']:
    el = OxmlElement(f'w:{side}')
    el.set(qn('w:val'), 'single')
    el.set(qn('w:sz'), '12')
    el.set(qn('w:space'), '18')
    el.set(qn('w:color'), '000000')
    pgBorders.append(el)
sectPr.append(pgBorders)

styles = doc.styles
styles['Normal'].font.name = 'Times New Roman'
styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
styles['Normal'].font.size = Pt(14)

def para(text='', size=14, bold=False, space_after=0, align=WD_ALIGN_PARAGRAPH.CENTER):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    r.font.size = Pt(size)
    r.bold = bold
    return p

para('HỌC VIỆN NÔNG NGHIỆP VIỆT NAM', 16, True, 0)
for _ in range(4): para('', 14, False, 0)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(28)
rn = p.add_run()
rn.add_picture(str(logo), width=Inches(1.55))

para('BÁO CÁO THỰC TẬP', 24, True, 4)
para('XÂY DỰNG WEBSITE QUẢN LÝ PHÒNG KHÁM THÚ Y', 18, True, 38)

# supervisor/student block with label/value layout like the sample
for label, value in [
    ('Giảng viên hướng dẫn:', 'ThS. NGUYỄN THỊ LAN'),
    ('Sinh viên thực hiện:', 'Trần Hoàng Minh    671688'),
    ('', 'Nguyễn Đức Anh    671'),
    ('', 'Nguyễn Quang Khải    671'),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    if label:
        r = p.add_run(label)
        r.bold = True
        r.font.name = 'Times New Roman'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        r.font.size = Pt(15)
        p.add_run('     ')
    else:
        p.add_run(' ' * 34)
    r2 = p.add_run(value)
    r2.bold = True
    r2.font.name = 'Times New Roman'
    r2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    r2.font.size = Pt(15)

for _ in range(6): para('', 14, False, 0)
para('HÀ NỘI - 2026', 14, True, 0)

doc.save(out)
print(out)
